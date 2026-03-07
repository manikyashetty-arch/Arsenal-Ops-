"""
PRD Processor Service - Extract text from PDF and Word documents
"""
import io
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document


class PRDProcessor:
    """Service to extract and process text from PRD documents"""
    
    def extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            docx_file = io.BytesIO(file_content)
            document = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from Word document: {str(e)}")
    
    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text based on file extension"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return self.extract_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return self.extract_from_docx(file_content)
        elif filename_lower.endswith('.doc'):
            # .doc files need different handling, fallback to docx parser
            try:
                return self.extract_from_docx(file_content)
            except:
                raise ValueError("Legacy .doc format not supported. Please convert to .docx")
        elif filename_lower.endswith('.txt'):
            return file_content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
        
        # Join with single newlines, paragraphs with double
        result = []
        prev_empty = False
        
        for line in cleaned_lines:
            if not line:
                if not prev_empty:
                    result.append('')
                prev_empty = True
            else:
                result.append(line)
                prev_empty = False
        
        return '\n'.join(result)
    
    def process_prd(self, file_content: bytes, filename: str) -> dict:
        """Process PRD file and return structured content"""
        raw_text = self.extract_text(file_content, filename)
        cleaned_text = self.clean_text(raw_text)
        
        return {
            "filename": filename,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "word_count": len(cleaned_text.split()),
            "char_count": len(cleaned_text)
        }


# Singleton instance
prd_processor = PRDProcessor()
